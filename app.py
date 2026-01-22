import streamlit as st
from datetime import date
import io
import re
import os
import time 
import requests
from PIL import Image


# --- FILE PATHING & DIAGRAM MAPPING ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(BASE_DIR, "diagrams")

# Static assets (ensure these files exist in the /diagrams folder)
AWS_PN_LOGO = os.path.join(ASSETS_DIR, "aws partner logo.jpg")
ONETURE_LOGO = os.path.join(ASSETS_DIR, "oneture logo1.jpg")
AWS_ADV_LOGO = os.path.join(ASSETS_DIR, "aws advanced logo1.jpg")

# Mapped Infra Costs
SOW_COST_TABLE_MAP = { 
    "L1 Support Bot POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Beauty Advisor POC SOW": { 
        "poc_cost": "4,525.66 USD + 200 USD (Amazon Bedrock Cost) = 4,725.66", 
        "prod_cost": "4,525.66 USD + 1,175.82 USD (Amazon Bedrock Cost) = 5,701.48" 
    }, 
    "Ready Search POC Scope of Work Document":{ "poc_cost": "2,641.40 USD" }, 
    "AI based Image Enhancement POC SOW": { "poc_cost": "2,814.34 USD" }, 
    "AI based Image Inspection POC SOW": { "poc_cost": "3,536.40 USD" }, 
    "Gen AI for SOP POC SOW": { "poc_cost": "2,110.30 USD" }, 
    "Project Scope Document": { "prod_cost": "2,993.60 USD" }, 
    "Gen AI Speech To Speech": { "prod_cost": "2,124.23 USD" }, 
    "PoC Scope Document": { "amazon_bedrock": "1,000 USD", "total": "$ 3,150" }
}

# Mapped Calculator Links
CALCULATOR_LINKS = {
    "L1 Support Bot POC SOW": "https://calculator.aws/#/estimate?id=211ea64cba5a8f5dc09805f4ad1a1e598ef5238b",
    "Ready Search POC Scope of Work Document": "https://calculator.aws/#/estimate?id=f8bc48f1ae566b8ea1241994328978e7e86d3490",
    "AI based Image Enhancement POC SOW": "https://calculator.aws/#/estimate?id=9a3e593b92b796acecf31a78aec17d7eb957d1e5",
    "Beauty Advisor POC SOW": "https://calculator.aws/#/estimate?id=3f89756a35f7bac7b2cd88d95f3e9aba9be9b0eb",
    "Beauty Advisor Production": "https://calculator.aws/#/estimate?id=4d7f092e819c799f680fd14f8de3f181f565c48e",
    "AI based Image Inspection POC SOW": "https://calculator.aws/#/estimate?id=72c56f93b0c0e101d67a46af4f4fe9886eb93342",
    "Gen AI for SOP POC SOW": "https://calculator.aws/#/estimate?id=c21e9b242964724bf83556cfeee821473bb935d1",
    "Project Scope Document": "https://calculator.aws/#/estimate?id=37339d6e34c73596559fe09ca16a0ac2ec4c4252",
    "Gen AI Speech To Speech": "https://calculator.aws/#/estimate?id=8444ae26e6d61e5a43e8e743578caa17fd7f3e69",
    "PoC Scope Document": "https://calculator.aws/#/estimate?id=420ed9df095e7824a144cb6c0e9db9e7ec3c4153"
}

SOW_DIAGRAM_MAP = {
    "L1 Support Bot POC SOW": os.path.join(ASSETS_DIR, "L1 Support Bot POC SOW.png"),
    "Ready Search POC Scope of Work Document": os.path.join(ASSETS_DIR, "Ready Search POC Scope of Work Document.png"),
    "AI based Image Enhancement POC SOW": os.path.join(ASSETS_DIR, "AI based Image Enhancement POC SOW.png"),
    "Beauty Advisor POC SOW": os.path.join(ASSETS_DIR, "Beauty Advisor POC SOW.png"),
    "AI based Image Inspection POC SOW": os.path.join(ASSETS_DIR, "AI based Image Inspection POC SOW.png"),
    "Gen AI for SOP POC SOW": os.path.join(ASSETS_DIR, "Gen AI for SOP POC SOW.png"),
    "Project Scope Document": os.path.join(ASSETS_DIR, "Project Scope Document.png"),
    "Gen AI Speech To Speech": os.path.join(ASSETS_DIR, "Gen AI Speech To Speech.png"),
    "PoC Scope Document": os.path.join(ASSETS_DIR, "PoC Scope Document.png")
}

# --- CONFIGURATION ---
st.set_page_config(
    page_title="GenAI SOW Architect", 
    layout="wide", 
    page_icon="📄",
    initial_sidebar_state="expanded"
)

# Custom CSS for an Enterprise UI
st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .stButton>button { border-radius: 8px; font-weight: 600; }
    .stTextArea textarea { border-radius: 10px; }
    .stTextInput input { border-radius: 8px; }
    .block-container { padding-top: 1.5rem; }
    .sow-preview {
        background-color: white;
        padding: 40px;
        border-radius: 12px;
        border: 1px solid #e2e8f0;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        line-height: 1.7;
        color: #1e293b;
        box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);
    }
    h1, h2, h3 { color: #0f172a; }
    .stTabs [data-baseweb="tab-list"] { gap: 24px; }
    .stTabs [data-baseweb="tab"] { height: 50px; white-space: pre-wrap; font-weight: 600; }
    [data-testid="stExpander"] { border: none; box-shadow: none; background: transparent; }
    .stakeholder-header { 
        background-color: #f1f5f9; 
        padding: 8px 12px; 
        border-radius: 6px; 
        margin-bottom: 10px; 
        font-weight: bold;
        color: #334155;
        border-left: 4px solid #3b82f6;
    }
    </style>
    """, unsafe_allow_html=True)

# Helper for docx hyperlinks
def add_hyperlink(paragraph, text, url):
    import docx.oxml.shared
    import docx.opc.constants
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_poc_calculation_table(doc):
    doc.add_paragraph("The above numbers are calculated basis the following:")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Particulars"
    hdr[1].text = "Value (in Dollar)"
    hdr[2].text = "Remarks"

    data = [
        ("Number of documents", "200", "Assuming 5 interactions for finalising each product copy"),
        ("Input Tokens per document", "10,00,000", ""),
        ("Input Token Cost per 1,000 Tokens", "0", "Anthropic Claude 3 Sonnet Model"),
        ("Total Input Cost in USD", "600", ""),
        ("Output Tokens per document", "50,000", ""),
        ("Output Token Cost per 1,000 Tokens", "0", "Anthropic Claude 3 Sonnet Model"),
        ("Total Output Cost in USD", "150", ""),
        ("Total Cost in USD", "750", ""),
        ("", "", ""),
        ("Tokens for Embedding Model", "2,50,00,00,000", ""),
        ("Input Cost per 1,000 Tokens", "0", "Cohere English Model"),
        ("Total Embedding Model Cost in USD", "250", ""),
        ("", "", ""),
        ("Total Cost in USD per month", "1,000", "")
    ]

    for row in data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


# WORD – COST TABLE (Section 5)
def add_infra_cost_table(doc, sow_type_name, text_content):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cost_data = SOW_COST_TABLE_MAP.get(sow_type_name)
    if not cost_data:
        return

    # Determine calculator link
    calc_url = CALCULATOR_LINKS.get(sow_type_name, "https://calculator.aws/#/")
    if sow_type_name == "Beauty Advisor POC SOW" and "Production Development" in text_content:
        calc_url = CALCULATOR_LINKS["Beauty Advisor Production"]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "System"
    hdr[1].text = "Infra Cost / month"
    hdr[2].text = "AWS Calculator Cost"

    rows_to_add = []
    if "poc_cost" in cost_data:
        rows_to_add.append(("POC", cost_data["poc_cost"]))
    if "prod_cost" in cost_data:
        rows_to_add.append(("Production", cost_data["prod_cost"]))
    if "amazon_bedrock" in cost_data:
        rows_to_add.append(("Amazon Bedrock", cost_data["amazon_bedrock"]))
    if "total" in cost_data:
        rows_to_add.append(("Total", cost_data["total"]))

    for label, cost in rows_to_add:
        r = table.add_row().cells
        r[0].text = label
        r[1].text = cost
        p = r[2].paragraphs[0]
        add_hyperlink(p, "Estimate", calc_url)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---- PoC Scope Document extra calculation table ----
    if sow_type_name == "PoC Scope Document":
        doc.add_paragraph("")  # spacing
        add_poc_calculation_table(doc)
        


        

# --- CACHED UTILITIES ---
def create_docx_logic(text_content, branding_info, sow_type_name):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ---- RIGID HEADER DEFINITIONS (CRITICAL) ----
    header_patterns = {
        "1": "1 TABLE OF CONTENTS",
        "2": "2 PROJECT OVERVIEW",
        "3": "3 ASSUMPTIONS & DEPENDENCIES",
        "4": "4 PROJECT SUCCESS CRITERIA",
        "5": "5 SCOPE OF WORK",
        "6": "6 SOLUTION ARCHITECTURE",
        "7": "7 PERFORMANCE & SECURITY",
        "8": "8 COST ESTIMATION",
        "9": "9 RESOURCES & COST ESTIMATES",
        "10": "10 Final Outputs"
    }



    architecture_rendered = False

    
    # State tracking to ensure rigid flow and prevent duplicates
    rendered_sections = {
        "1": False,  # TOC
        "2": False,  # Project Overview
        "3": False,  # Assumptions & Dependencies
        "4": False,  # Project Success Criteria
        "5": False,  # Scope of Work
        "6": False,  # Solution Architecture
        "7": False,  # Performance & Security
        "8": False,  # Cost Estimation
        "9": False   # Resources & Cost Estimates
    }


    # --- PAGE 1: COVER PAGE ---
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(AWS_PN_LOGO):
        doc.add_picture(AWS_PN_LOGO, width=Inches(1.6))

    doc.add_paragraph("\n" * 3)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(branding_info['sow_name'])
    run.font.size = Pt(26)
    run.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.add_run("Scope of Work Document").font.size = Pt(14)

    doc.add_paragraph("\n" * 4)

    logo_table = doc.add_table(rows=1, cols=3)
    logo_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Customer Logo
    cell = logo_table.rows[0].cells[0]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if branding_info.get("customer_logo_bytes"):
        cell.paragraphs[0].add_run().add_picture(io.BytesIO(branding_info["customer_logo_bytes"]), width=Inches(1.8))
    else:
        cell.paragraphs[0].add_run("Customer Logo").bold = True

    # Oneture Logo
    cell = logo_table.rows[0].cells[1]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(ONETURE_LOGO):
        cell.paragraphs[0].add_run().add_picture(ONETURE_LOGO, width=Inches(2.2))

    # AWS Advanced Tier
    cell = logo_table.rows[0].cells[2]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(AWS_ADV_LOGO):
        cell.paragraphs[0].add_run().add_picture(AWS_ADV_LOGO, width=Inches(1.8))

    doc.add_paragraph("\n" * 3)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(branding_info["doc_date_str"]).bold = True

    doc.add_page_break()

    # --- CONTENT PROCESSING ---
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lines = text_content.split('\n')
    i = 0
    in_toc_section = False
    content_started = False

    # Define Header Rigid Flow
   




    while i < len(lines):
        line = lines[i].strip()
        if not line:
            if i > 0 and lines[i-1].strip() and content_started: doc.add_paragraph("")
            i += 1
            continue

        line_clean = re.sub(r'\*+', '', line).strip()
        clean_text = re.sub(r'^#+\s*', '', line_clean).strip()
        upper_text = clean_text.upper()

        # Check if line matches a main section trigger
        current_header_id = None
        for h_id, pattern in header_patterns.items():
            if upper_text.startswith(pattern):
                current_header_id = h_id
                break

        # Remove unnecessary commentary, triggers, and redundant AI descriptions
        irrelevant_keywords = ["PLACEHOLDER FOR COST TABLE", "SPECIFICS TO BE DISCUSSED BASIS POC"]
        if any(kw in upper_text for kw in irrelevant_keywords):
            i += 1
            continue

        # Content started guard: Skip introductory fluff
        if not content_started:
            if current_header_id == "1":
                content_started = True
            else:
                i += 1
                continue

        # Handle Section Switches (Enforcing Single Rendering)
        if current_header_id:
            # Enforce Page breaks for TOC (Page 2) and Overview (Page 3)
            if in_toc_section and current_header_id == "2":
                in_toc_section = False
                doc.add_page_break()
            
            if current_header_id:
                if current_header_id not in rendered_sections:
                    i += 1
                    continue

                if not rendered_sections[current_header_id]:
                    doc.add_heading(clean_text, level=1)
                    rendered_sections[current_header_id] = True


                # Immediate content injection
                # Handle Section Switches (Enforcing Single Rendering)

                if current_header_id == "6" and not architecture_rendered:
                    architecture_rendered = True

                    diagram_path = SOW_DIAGRAM_MAP.get(sow_type_name)

                    if diagram_path and os.path.exists(diagram_path):
                        try:
                            doc.add_paragraph("")
                            doc.add_picture(diagram_path, width=Inches(6.0))
                            cap = doc.add_paragraph(f"{sow_type_name} – Architecture Diagram")
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception:
                            doc.add_paragraph("Architecture diagram could not be rendered.")
                else:
                    doc.add_paragraph("Architecture diagram not available for this use case.")

                i += 1
                continue


                if current_header_id == "5" or "COST ESTIMATION" in upper_text:
                    add_infra_cost_table(doc, sow_type_name, text_content)

            
            i += 1
            continue

        # ---------------- TABLE PARSING ----------------
        if line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            # Filter out redundant tables generated for Section 5
            if rendered_sections["5"] and not rendered_sections["6"]:
                 i += 1
                 continue

            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 3:
                headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for idx, h in enumerate(headers): table.rows[0].cells[idx].text = h
                for row_line in table_lines[2:]:
                    row_cells = table.add_row().cells
                    cells = [c.strip() for c in row_line.split('|') if c.strip()]
                    for idx, c in enumerate(cells):
                        if idx < len(row_cells): row_cells[idx].text = c
            continue

        # ---------------- HEADINGS (Levels 2 and 3) ----------------
        if line.startswith('## '):
            h = doc.add_heading(clean_text, level=2)
            if in_toc_section: h.paragraph_format.left_indent = Inches(0.4)
        elif line.startswith('### '):
            h = doc.add_heading(clean_text, level=3)
            if in_toc_section: h.paragraph_format.left_indent = Inches(0.8)
        
        # ---------------- BULLETS ----------------
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(clean_text[2:] if (clean_text.startswith('- ') or clean_text.startswith('* ')) else clean_text, style="List Bullet")
            if in_toc_section: p.paragraph_format.left_indent = Inches(0.4)
        
        # ---------------- NORMAL TEXT ----------------
        else:
            # Skip architectural descriptions that AI adds which repeat diagram info
            
            
            p = doc.add_paragraph(clean_text)
            bold_keywords = [
                "PARTNER EXECUTIVE SPONSOR", "CUSTOMER EXECUTIVE SPONSOR", 
                "AWS EXECUTIVE SPONSOR", "PROJECT ESCALATION CONTACTS", 
                "ASSUMPTIONS:", "DEPENDENCIES:", "ASSUMPTIONS (", "DEPENDENCIES ("
            ]
            if any(k in upper_text for k in bold_keywords):
                if p.runs: p.runs[0].bold = True
        i += 1
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# API CALL WRAPPER WITH RETRY LOGIC (Exponential Backoff)
def call_gemini_with_retry(api_key, payload):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    retries = 5
    for attempt in range(retries):
        try:
            res = requests.post(url, json=payload)
            if res.status_code == 200:
                return res, None
            # If 503 (Overloaded) or other transient errors, wait and retry
            if res.status_code in [503, 429]:
                time.sleep(2**attempt)
                continue
            else:
                return None, f"API Error {res.status_code}: {res.text}"
        except Exception as e:
            time.sleep(2**attempt)
    return None, "The model is currently overloaded after multiple retries. Please try again in a few moments."

# --- INITIALIZATION ---
# --- SESSION STATE INITIALIZATION ---
if "engagement_type" not in st.session_state:
    st.session_state.engagement_type = "Proof of Concept (PoC)"

if "success_dimensions" not in st.session_state:
    st.session_state.success_dimensions = []

if "customer_dependencies" not in st.session_state:
    st.session_state.customer_dependencies = []

if "data_types" not in st.session_state:
    st.session_state.data_types = []

if "data_characteristics" not in st.session_state:
    st.session_state.data_characteristics = {}

if "data_types" not in st.session_state:
    st.session_state.data_types = []

if "key_assumptions" not in st.session_state:
    st.session_state.key_assumptions = []

if "other_assumptions" not in st.session_state:
    st.session_state.other_assumptions = ""

if "user_validation_required" not in st.session_state:
    st.session_state.user_validation_required = "Yes – customer validation required"

if "compute_orchestration" not in st.session_state:
    st.session_state.compute_orchestration = "AWS Lambda + Step Functions"

if "genai_services" not in st.session_state:
    st.session_state.genai_services = ["Amazon Bedrock (LLM inference)"]

if "storage_services" not in st.session_state:
    st.session_state.storage_services = ["Amazon S3"]

if "ui_layer" not in st.session_state:
    st.session_state.ui_layer = "Streamlit on S3"

if "performance_expectation" not in st.session_state:
    st.session_state.performance_expectation = "Batch"

if "security_compliance" not in st.session_state:
    st.session_state.security_compliance = []

if "cost_ownership" not in st.session_state:
    st.session_state.cost_ownership = "Funded by Customer"

if "deliverables" not in st.session_state:
    st.session_state.deliverables = []

if "post_poc_next_steps" not in st.session_state:
    st.session_state.post_poc_next_steps = []

if "poc_duration" not in st.session_state:
    st.session_state.poc_duration = "4 weeks"

if "phase_breakdown" not in st.session_state:
    st.session_state.phase_breakdown = {
        "Infra setup": "",
        "Core workflows": "",
        "Testing & validation": "",
        "Demo & feedback": ""
    }











if 'generated_sow' not in st.session_state:
    st.session_state.generated_sow = ""

if 'stakeholders' not in st.session_state:
    import pandas as pd
    st.session_state.stakeholders = {
        "Partner": pd.DataFrame([{"Name": "Gaurav Kankaria", "Title": "Head of Analytics & ML", "Email": "gaurav.kankaria@oneture.com"}]),
        "Customer": pd.DataFrame([{"Name": "Cheten Dev", "Title": "Head of Product Design", "Email": "cheten.dev@nykaa.com"}]),
        "AWS": pd.DataFrame([{"Name": "Anubhav Sood", "Title": "AWS Account Executive", "Email": "anbhsood@amazon.com"}]),
        "Escalation": pd.DataFrame([
            {"Name": "Omkar Dhavalikar", "Title": "AI/ML Lead", "Email": "omkar.dhavalikar@oneture.com"},
            {"Name": "Gaurav Kankaria", "Title": "Head of Analytics and AIML", "Email": "gaurav.kankaria@oneture.com"}
        ])
    }

def clear_sow():
    st.session_state.generated_sow = ""

# --- SIDEBAR: PROJECT INTAKE ---
with st.sidebar:
    st.title("SOW Architect")
    
    with st.expander("API Key", expanded=False):
        api_key = st.text_input("Gemini API Key", type="password")
    
    st.divider()
    st.header(" 1. Project Intake")
    st.subheader("1.2 Engagement Type")

    st.radio(
         "Select engagement type:",
         [
             "Proof of Concept (PoC)",
             "Pilot",
             "MVP",
             "Production Rollout",
             "Assessment / Discovery",
             "Support"
         ],
         key="engagement_type"
     )



    sow_type_options = list(SOW_COST_TABLE_MAP.keys())
    selected_sow_name = st.selectbox("Scope of Work Type", sow_type_options)

    st.divider()
    industry_options = ["Retail / E-commerce", "BFSI", "Manufacturing", "Telecom", "Healthcare", "Energy / Utilities", "Logistics", "Media", "Government", "Other (specify)"]
    industry_type = st.selectbox("Industry / Domain", industry_options)
    final_industry = st.text_input("Specify Industry", placeholder="Enter industry...") if industry_type == "Other (specify)" else industry_type
    duration = st.text_input("Timeline / Duration", "4 Weeks")
    
    if st.button(" Reset All Fields", on_click=clear_sow, use_container_width=True):
        st.rerun()

# --- MAIN UI ---
st.title(" GenAI Scope of Work Architect")
st.header(" Cover Page Branding")
customer_logo = st.file_uploader("Upload Customer Logo (Optional)", type=["png", "jpg", "jpeg"])
doc_date = st.date_input("Document Date", date.today())
st.divider()

st.header("2. Objectives & Stakeholders")
st.subheader(" 2.1 Objective")
objective = st.text_area("Define the core business objective:", placeholder="e.g., Development of a Gen AI based WIMO Bot...", height=120)
outcomes = st.multiselect("Key Outcomes:", ["Reduce manual effort", "Improve accuracy / quality", "Faster turnaround time", "Cost reduction", "Revenue uplift", "Compliance improvement", "Better customer experience", "Scalability validation", "Integration Feasibility"])
st.divider()

st.subheader(" 2.3 Project Sponsor(s) / Stakeholder(s) / Project Team")
col_team1, col_team2 = st.columns(2)
with col_team1:
    st.markdown('<div class="stakeholder-header">Partner Executive Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Partner"] = st.data_editor(st.session_state.stakeholders["Partner"], num_rows="dynamic", use_container_width=True, key="ed_partner")
    st.markdown('<div class="stakeholder-header">AWS Executive Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["AWS"] = st.data_editor(st.session_state.stakeholders["AWS"], num_rows="dynamic", use_container_width=True, key="ed_aws")
with col_team2:
    st.markdown('<div class="stakeholder-header">Customer Executive Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Customer"] = st.data_editor(st.session_state.stakeholders["Customer"], num_rows="dynamic", use_container_width=True, key="ed_customer")
    st.markdown('<div class="stakeholder-header">Project Escalation Contacts</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Escalation"] = st.data_editor(st.session_state.stakeholders["Escalation"], num_rows="dynamic", use_container_width=True, key="ed_escalation")

# --- 3.1 Customer Dependencies ---
st.divider()
st.header("3.1 Customer Dependencies")

dependencies = st.multiselect(
    "Select all that apply:",
    [
        "Sample data availability",
        "Historical data availability",
        "Design / business guidelines finalized",
        "API access provided",
        "User access to AWS account",
        "SME availability for validation",
        "Network / VPC access",
        "Security approvals"
    ],
    key="customer_dependencies"
)

# --- 3.2 Data Characteristics ---
st.divider()
st.header("3.2 Data Characteristics")

data_types = st.multiselect(
    "What type of data is involved?",
    [
        "Images",
        "Text",
        "PDFs / Documents",
        "Audio",
        "Video",
        "Structured tables",
        "APIs / Streams"
    ],
    key="data_types"
)

data_characteristics = {}

st.subheader("3.3 Key Assumptions")

assumption_options = [
    "PoC only, not production-grade",
    "Limited data volume",
    "Rule-based logic acceptable initially",
    "Manual review for edge cases",
    "No real-time SLA commitments"
]

selected_assumptions = st.multiselect(
    "Select applicable assumptions:",
    assumption_options,
    default=st.session_state.key_assumptions
)

other_assumption_text = st.text_area(
    "Other assumptions (optional):",
    value=st.session_state.other_assumptions,
    placeholder="Enter any additional assumptions..."
)

# Store safely in session_state
st.session_state.key_assumptions = selected_assumptions
st.session_state.other_assumptions = other_assumption_text

st.subheader("4.1 Success Dimensions")

success_options = [
    "Accuracy",
    "Latency",
    "Usability",
    "Explainability",
    "Coverage",
    "Cost efficiency",
    "Integration readiness"
]

selected_success_dimensions = st.multiselect(
    "Select success dimensions:",
    success_options,
    default=st.session_state.success_dimensions
)

st.session_state.success_dimensions = selected_success_dimensions

st.subheader("4.2 User Validation Requirement")

st.radio(
    "Select validation approach:",
    [
        "Yes – customer validation required",
        "No – internal validation sufficient"
    ],
    key="user_validation_required"
)

st.subheader("6.1 Compute & Orchestration")

st.radio(
    "Select compute & orchestration approach:",
    [
        "AWS Lambda",
        "Step Functions",
        "AWS Lambda + Step Functions",
        "ECS / EKS (future)",
        "Hybrid"
    ],
    key="compute_orchestration"
)

st.subheader("6.2 GenAI / ML Services")

st.multiselect(
    "Select GenAI / ML services to be used:",
    [
        "Amazon Bedrock (LLM inference)",
        "SageMaker (custom models)",
        "Rekognition",
        "Textract",
        "Comprehend",
        "Transcribe",
        "Translate"
    ],
    default=st.session_state.genai_services,
    key="genai_services"
)

st.subheader("6.3 Storage & Search")

st.multiselect(
    "Select storage and search services:",
    [
        "Amazon S3",
        "DynamoDB",
        "OpenSearch",
        "RDS",
        "Vector DB (OpenSearch / Aurora PG)"
    ],
    default=st.session_state.storage_services,
    key="storage_services"
)

st.subheader("6.4 UI Layer")

st.radio(
    "Select UI deployment option:",
    [
        "Streamlit on S3",
        "CloudFront + Static UI",
        "Internal demo only",
        "No UI (API only)"
    ],
    index=[
        "Streamlit on S3",
        "CloudFront + Static UI",
        "Internal demo only",
        "No UI (API only)"
    ].index(st.session_state.ui_layer),
    key="ui_layer"
)

st.subheader("7.1 Performance Expectations")

st.selectbox(
    "Select performance expectation:",
    [
        "Batch",
        "Near real-time",
        "Real-time"
    ],
    index=[
        "Batch",
        "Near real-time",
        "Real-time"
    ].index(st.session_state.performance_expectation),
    key="performance_expectation"
)

st.subheader("7.2 Security & Compliance")

st.multiselect(
    "Select applicable security and compliance requirements:",
    [
        "IAM-based access",
        "Encryption at rest",
        "Encryption in transit",
        "VPC deployment",
        "Audit logging",
        "Compliance alignment (RBI, SOC2, etc.)"
    ],
    default=st.session_state.security_compliance,
    key="security_compliance"
)

st.subheader("8. Timeline & Phasing")

st.markdown("### 8.1 PoC Duration")
st.session_state.poc_duration = st.radio(
    "Select PoC duration:",
    ["2 weeks", "4 weeks", "6 weeks", "Custom"]
)

st.markdown("### 8.2 Phase Breakdown")

st.info("You can optionally map weeks to each phase")

for phase in st.session_state.phase_breakdown:
    st.session_state.phase_breakdown[phase] = st.text_input(
        f"{phase} (e.g., Week 1–2)",
        value=st.session_state.phase_breakdown[phase]
    )


st.subheader("9 Cost Ownership")

st.radio(
    "Select cost ownership model:",
    [
        "Funded by AWS",
        "Funded by Partner",
        "Funded by Customer",
        "Shared"
    ],
    index=[
        "Funded by AWS",
        "Funded by Partner",
        "Funded by Customer",
        "Shared"
    ].index(st.session_state.cost_ownership),
    key="cost_ownership"
)

st.subheader("10. Final Outputs")

st.markdown("### 10.1 Deliverables")
st.session_state.deliverables = st.multiselect(
    "Select deliverables:",
    [
        "PoC architecture",
        "Working demo",
        "SOW document",
        "Cost estimate",
        "Next-phase proposal"
    ]
)

st.markdown("### 10.2 Post-PoC Next Steps")
st.session_state.post_poc_next_steps = st.multiselect(
    "Select post-PoC next steps:",
    [
        "Production proposal",
        "Scaling roadmap",
        "Security review",
        "Performance optimization",
        "Model fine-tuning"
    ]
)










if "Images" in data_types:
    st.subheader("Images")
    data_characteristics["Images"] = {
        "avg_size_mb": st.text_input("Average image size (MB)", key="img_size"),
        "formats": st.text_input("Formats (JPEG, PNG, etc.)", key="img_formats"),
        "volume": st.text_input("Approx volume (per day / total)", key="img_volume"),
    }

if "Text" in data_types:
    st.subheader("Text")
    data_characteristics["Text"] = {
        "source": st.text_input("Source (chat, logs, docs)", key="txt_source"),
        "volume": st.text_input("Approx volume", key="txt_volume"),
    }

if "PDFs / Documents" in data_types:
    st.subheader("PDFs / Documents")
    data_characteristics["Documents"] = {
        "avg_pages": st.text_input("Average pages per document", key="pdf_pages"),
        "volume": st.text_input("Approx volume", key="pdf_volume"),
    }

st.session_state.data_characteristics = data_characteristics



# --- GENERATION ---
if st.button("✨ Generate SOW Document", type="primary", use_container_width=True):
    if not api_key:
        st.warning("⚠️ Enter a Gemini API Key in the sidebar.")
    elif not objective:
        st.error("⚠️ Business Objective is required.")
    else:
        with st.spinner(f"Architecting {selected_sow_name}..."):
            def get_md(df): return df.to_markdown(index=False)
            
            # Dynamic Table Context
            cost_info = SOW_COST_TABLE_MAP.get(selected_sow_name, {})
            dynamic_table_prompt = "| System | Infra Cost / month | AWS Calculator Cost |\n| --- | --- | --- |\n"
            if "poc_cost" in cost_info:
                dynamic_table_prompt += f"| POC | {cost_info['poc_cost']} | Estimate |\n"
            if "prod_cost" in cost_info:
                dynamic_table_prompt += f"| Production | {cost_info['prod_cost']} | Estimate |\n"
            if "amazon_bedrock" in cost_info:
                dynamic_table_prompt += f"| Amazon Bedrock | {cost_info['amazon_bedrock']} | Estimate |\n"
            if "total" in cost_info:
                dynamic_table_prompt += f"| Total | {cost_info['total']} | Estimate |\n"

            prompt_text = f"""
            Generate a COMPLETE formal enterprise SOW for {selected_sow_name} in {final_industry}.

            You are given a FIXED SOW STRUCTURE.
            You must ONLY generate CONTENT under each section.
            DO NOT create or rename section headings.
            DO NOT reorder sections.
            DO NOT add new sections.

            COST RULES (CRITICAL):
            - Do NOT generate cost tables.
            - Do NOT mention dollar values.
            - Cost section must be narrative only.


            STRUCTURE TEMPLATE (DO NOT MODIFY):

            1 TABLE OF CONTENTS
            (leave blank)

            2 PROJECT OVERVIEW
            2.1 OBJECTIVE
            2.2 PROJECT TEAM

            3 ASSUMPTIONS & DEPENDENCIES
            3.1 Customer Dependencies
            3.2 Data Characteristics
            3.3 Key Assumptions

            4 PROJECT SUCCESS CRITERIA
            4.1 Success Dimensions
            4.2 User Validation Requirement

            5 SCOPE OF WORK - TECHNICAL PROJECT PLAN
            (Fill detailed technical tasks)

            6 SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM
            (Write ONLY 3–5 bullet points. No diagram description.)

            7 PERFORMANCE & SECURITY
            7.1 Performance Expectations
            7.2 Security & Compliance

            8 COST ESTIMATION

            9 RESOURCES & COST ESTIMATES



            ENGAGEMENT CONTEXT:
            - Engagement Type: {st.session_state.engagement_type}
            - Adjust scope depth, success criteria strictness, assumptions, and cost modeling based on the engagement type.
            
            STRICT SECTION FLOW (OUTPUT EACH SECTION ONCE, NO REPETITION):
            1 TABLE OF CONTENTS
            2 PROJECT OVERVIEW
              2.1 OBJECTIVE: {objective}
              2.2 PROJECT TEAM:
                  ### Partner Executive Sponsor
                  {get_md(st.session_state.stakeholders["Partner"])}
                  ### Customer Executive Sponsor
                  {get_md(st.session_state.stakeholders["Customer"])}
                  ### AWS Executive Sponsor
                  {get_md(st.session_state.stakeholders["AWS"])}
                  ### Project Escalation Contacts
                  {get_md(st.session_state.stakeholders["Escalation"])}

              2.4 Project Success Criteria
            3.1 CUSTOMER DEPENDENCIES
            Selected by user:
            {", ".join(st.session_state.customer_dependencies) if st.session_state.customer_dependencies else "No explicit customer dependencies specified."}

            Instruction:
            - Expand each selected dependency into a formal enterprise dependency statement.
            - Clearly mention customer responsibility and prerequisite nature.

            3.2 DATA CHARACTERISTICS
            Selected data types:
            {", ".join(st.session_state.data_types) if st.session_state.data_types else "No data types specified."}

            Detailed inputs:
            {st.session_state.data_characteristics if st.session_state.data_characteristics else "No detailed data characteristics provided."}

            Instruction:
            - Use this information to influence architecture decisions.
            - Use this to justify Amazon Bedrock / ML service selection.
            - Reflect data volume, format, and frequency assumptions in cost rationale.

            3.3 Key Assumptions:
            Selected assumptions by user:
            {", ".join(st.session_state.key_assumptions) if st.session_state.key_assumptions else "No predefined assumptions selected."}

            Additional assumptions:
            {st.session_state.other_assumptions if st.session_state.other_assumptions.strip() else "None provided."}

            Instructions:
            - Convert each selected assumption into a formal, professional SOW assumption.
            - Do NOT repeat checkbox text verbatim.
            - Align assumptions with engagement type: {st.session_state.engagement_type}

            4.1 Project Success Criteria

            Selected success dimensions:
            {", ".join(st.session_state.success_dimensions) if st.session_state.success_dimensions else "No explicit success dimensions selected."}

            Instructions:
            - Generate measurable, quantifiable success criteria for EACH selected dimension.
            - Tailor criteria to the solution type: {selected_sow_name}.
            - Use realistic enterprise metrics.

            Examples (if applicable):
            - Accuracy → "≥85% match with manual reviewer outcomes"
            - Latency → "Average response time under 2 seconds"
            - Cost efficiency → "Operate within defined monthly inference budget"

            4.2 User Validation Requirement

            Validation approach selected:
            {st.session_state.user_validation_required}

            Instructions:
            - If customer validation is required, clearly state customer responsibilities and sign-off expectations.
            - If internal validation is sufficient, specify internal review and acceptance criteria.
            - Align validation approach with engagement type: {st.session_state.engagement_type}.
            
            5 SCOPE OF WORK - TECHNICAL PROJECT PLAN
            
            6 SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM

            6.1 Compute & Orchestration

            Selected compute approach:
            {st.session_state.compute_orchestration}

            Instructions:
            - Reflect this choice in the solution architecture narrative.
            - If ECS / EKS is selected, mention it as a future-state scalability option.
            - Align compute choice with engagement type: {st.session_state.engagement_type}.

            6.2 GenAI / ML Services

            Selected services:
            {", ".join(st.session_state.genai_services) if st.session_state.genai_services else "None selected"}

            Instructions:
            - Justify each selected service in the architecture.
            - Explicitly explain why Amazon Bedrock is chosen if selected.
            - Use selected services to influence pricing assumptions.
            - Align services with engagement type: {st.session_state.engagement_type}.

            6.3 Storage & Search

            Selected services:
            {", ".join(st.session_state.storage_services) if st.session_state.storage_services else "None selected"}

            Instructions:
            - Justify each selected storage or search service.
            - Clearly explain data flow between GenAI services and storage.
            - If Vector DB is selected, explain embedding storage and retrieval.
            - Reflect storage choices in cost and scalability considerations.

            6.4 UI Layer

            Selected UI approach:
            {st.session_state.ui_layer}

            Instructions:
            - Explain why this UI option is appropriate for the engagement type.
            - Mention hosting, access control, and demo expectations.
            - If "No UI (API only)" is selected, explicitly state API-only consumption.
            - Align UI choice with security and scalability assumptions.

            7 PERFORMANCE & SECURITY

            7.1 Performance Expectations:
            {st.session_state.performance_expectation}

            Instructions:
            - Explain processing model (batch vs real-time).
            - Align performance with architecture and service choices.
            - Mention latency expectations clearly.

            7.2 Security & Compliance:
            {", ".join(st.session_state.security_compliance) if st.session_state.security_compliance else "Standard AWS security best practices"}

            Instructions:
            - Expand selected items into formal enterprise security controls.
            - Clearly mention customer and partner responsibilities.
            - If compliance standards are selected, reflect governance alignment.

            8 TIMELINE & PHASING

            8.1 PoC Duration
            Selected duration:
            {st.session_state.poc_duration}

            Instruction:
            - Describe overall engagement duration.
            - Align depth of activities with selected duration.

            8.2 Phase Breakdown
            Phase mapping provided by user:
            {st.session_state.phase_breakdown}

            Instruction:
            - Expand each phase into a clear timeline narrative.
            - Ensure logical sequencing.
            - Keep enterprise consulting tone.


            9 Cost Ownership:
            {st.session_state.cost_ownership}

            Instructions:
            - Clearly state who bears infrastructure and GenAI service costs.
            - Reflect this ownership consistently in cost assumptions.
            - If "Shared", clearly explain cost split responsibility.

            10 FINAL OUTPUTS

            10.1 Deliverables
            Selected by user:
            {", ".join(st.session_state.deliverables) if st.session_state.deliverables else "Standard PoC deliverables"}

            Instruction:
            - Expand each deliverable into a professional outcome-oriented statement.
            - Clearly state what the customer receives.

            10.2 Post-PoC Next Steps
            Selected by user:
            {", ".join(st.session_state.post_poc_next_steps) if st.session_state.post_poc_next_steps else "To be mutually agreed"}

            Instruction:
            - Describe each next step as a logical progression after PoC.
                - Align with enterprise delivery best practices.



            
            """

            payload = {
                "contents": [{"parts": [{"text": prompt_text}]}],
                "systemInstruction": {"parts": [{"text": "Solutions Architect. Follow numbering exactly. Page 1 cover, Page 2 TOC, Page 3 starts Overview. No repetitions. No introductory fluff."}]}
            }
            

            res, error = call_gemini_with_retry(api_key, payload)
            if res:
                response_json = res.json()

                try:
                    candidates = response_json.get("candidates", [])
                    if not candidates:
                        raise ValueError("No candidates returned by Gemini")

                    content = candidates[0].get("content", {})
                    parts = content.get("parts", [])

                    text_parts = [p.get("text", "") for p in parts if "text" in p]
                    if not text_parts:
                        raise ValueError("No text content returned by Gemini")

                    st.session_state.generated_sow = "\n".join(text_parts)
                    st.balloons()

                except Exception as e:
                    st.error("❌ Failed to generate SOW content from Gemini.")
                    st.error(str(e))
                    st.json(response_json)

                st.balloons()
            else:
                st.error(error)

# --- STEP 3: REVIEW & EXPORT ---
if st.session_state.generated_sow:
    st.divider()
    st.header("3. Review & Export")
    tab_edit, tab_preview = st.tabs(["✍️ Document Editor", "📄 Visual Preview"])
    with tab_edit:
        st.session_state.generated_sow = st.text_area(label="Modify content:", value=st.session_state.generated_sow, height=700, key="sow_editor")
    with tab_preview:
        st.markdown(f'<div class="sow-preview">', unsafe_allow_html=True)
        # Handle links in preview
        calc_url_p = CALCULATOR_LINKS.get(selected_sow_name, "https://calculator.aws")
        if selected_sow_name == "Beauty Advisor POC SOW" and "Production Development" in st.session_state.generated_sow:
            calc_url_p = CALCULATOR_LINKS["Beauty Advisor Production"]
        preview_content = st.session_state.generated_sow.replace("Estimate", f'<a href="{calc_url_p}" target="_blank" style="color:#3b82f6; text-decoration: underline;">Estimate</a>')
        
        header_pattern = r'(?i)(^#*\s*\d+\s+SOLUTION ARCHITECTURE.*)'

        match = re.search(header_pattern, preview_content, re.MULTILINE)
        if match:
            start, end = match.span()
            st.markdown(preview_content, unsafe_allow_html=True)

            diagram_path_out = SOW_DIAGRAM_MAP.get(selected_sow_name)
            if diagram_path_out and os.path.exists(diagram_path_out):
                try:
                    img = Image.open(diagram_path_out)
                    st.image(
                        img,
                        caption=f"{selected_sow_name} – Architecture Diagram",
                        use_container_width=True
                    )
                except Exception as e:
                    st.warning("⚠️ Architecture diagram exists but is not a valid image file.")
                    st.caption(f"File path: {diagram_path_out}")
        else:
            st.info("ℹ️ No architecture diagram mapped for this use case.")

    if "customer_dependencies" not in st.session_state:
        st.session_state.customer_dependencies = []

    if "data_types" not in st.session_state:
        st.session_state.data_types = []

    if "data_characteristics" not in st.session_state:
        st.session_state.data_characteristics = {}

    if "success_dimensions" not in st.session_state:
        st.session_state.success_dimensions = []



    

    if st.button("💾 Prepare Microsoft Word Document"):
        branding_info = {
            "sow_name": selected_sow_name,
            "customer_logo_bytes": customer_logo.getvalue() if customer_logo else None,
            "doc_date_str": doc_date.strftime("%d %B %Y")
        }
        docx_data = create_docx_logic(st.session_state.generated_sow, branding_info, selected_sow_name)
        st.download_button(label="📥 Download Now (.docx)", data=docx_data, file_name=f"SOW_{selected_sow_name.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
